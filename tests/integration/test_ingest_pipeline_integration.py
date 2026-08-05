"""End-to-end integration for the RIT-I-0014 gated ingest pipeline.

Exercises the deterministic rails through the real CLI transport, proving the
redesigned ingest boundary works as a chain:

    init -> extract-text -> (faithful ResumeDocument) -> validate-faithfulness
         -> set-active -> downstream check-ats-structure

plus the negative path (a lossy conversion is rejected with a non-zero exit),
the job pointer path, and deterministic docx extraction (no optional extras).

All assertions are deterministic: no LLM, no network. The happy-path fixture is
crafted against the finalized faithfulness rules (RIT-T-0092) — every JSON token
appears in the source, high-signal field values appear verbatim/contiguous in
the source, and source/JSON bullet counts match — so the gate passes with at
most heuristic warnings (never an error-severity finding).
"""

from __future__ import annotations

import importlib.util
import json
from pathlib import Path

from resume_kit_cli.app import app
from resume_kit_schemas import ResumeDocument
from typer.testing import CliRunner

runner = CliRunner()


def _docx_available() -> bool:
    """Return True when python-docx is importable (a base-install dependency)."""
    return importlib.util.find_spec("docx") is not None


# Raw source text whose tokens are a superset of the faithful ResumeDocument's,
# with every high-signal field value (company/title/dates, institution/degree/
# years) appearing as a contiguous run, and exactly two bullet-prefixed lines.
_SOURCE_TEXT = (
    "Ada Lovelace\n"
    "Senior Engineer\n"
    "Engineer, Acme Corp, 2020-2024\n"
    "- Built Python services for analytics\n"
    "- Led data pipeline migration\n"
    "Python, SQL\n"
    "BSc Computer Science, MIT, 2018\n"
)


def _faithful_resume() -> ResumeDocument:
    return ResumeDocument.model_validate(
        {
            "personalInfo": {"name": "Ada Lovelace", "title": "Senior Engineer"},
            "workExperience": [
                {
                    "id": 1,
                    "title": "Engineer",
                    "company": "Acme Corp",
                    "years": "2020-2024",
                    "description": [
                        "Built Python services for analytics",
                        "Led data pipeline migration",
                    ],
                }
            ],
            "education": [
                {
                    "id": 1,
                    "institution": "MIT",
                    "degree": "BSc Computer Science",
                    "years": "2018",
                }
            ],
            "additional": {"technicalSkills": ["Python", "SQL"]},
        }
    )


def _init_project(root: Path) -> None:
    result = runner.invoke(app, ["init", "--root", str(root)])
    assert result.exit_code == 0, result.stdout
    assert (root / "resume-kit" / "config.json").is_file()
    for sub in ("resumes", "jobs", "working", "learning"):
        assert (root / "resume-kit" / sub).is_dir()


def test_happy_path_ingest_chain(tmp_path: Path) -> None:
    _init_project(tmp_path)
    rk = tmp_path / "resume-kit"

    # 1) deterministic extraction of the source text (identity for .txt)
    src = tmp_path / "resume-a.txt"
    src.write_text(_SOURCE_TEXT, encoding="utf-8")
    extracted = runner.invoke(app, ["extract-text", str(src)])
    assert extracted.exit_code == 0, extracted.stdout
    payload = json.loads(extracted.stdout)
    assert payload["errors"] == []
    assert "Ada Lovelace" in payload["data"]["text"]

    # 2) a faithful ResumeDocument passes the HARD GATE (exit 0, passed True)
    resume_json = rk / "resumes" / "resume-a-original.json"
    resume_json.write_text(_faithful_resume().model_dump_json(), encoding="utf-8")
    gate = runner.invoke(
        app, ["validate-faithfulness", "--source", str(src), "--json", str(resume_json)]
    )
    assert gate.exit_code == 0, gate.stdout
    report = json.loads(gate.stdout)["data"]
    assert report["passed"] is True
    assert [f for f in report["findings"] if f["severity"] == "error"] == []

    # 3) set-active records the pointer + source path through the code-owned schema
    activate = runner.invoke(
        app,
        [
            "set-active",
            "--root",
            str(tmp_path),
            "--resume",
            "resumes/resume-a-original.json",
            "--source",
            str(src),
        ],
    )
    assert activate.exit_code == 0, activate.stdout
    config = json.loads((rk / "config.json").read_text(encoding="utf-8"))
    assert config["active_resume"] == "resumes/resume-a-original.json"
    assert config["active_resume_source"] == str(src)

    # 4) downstream deterministic analysis still runs clean on the ingested resume
    structure = runner.invoke(
        app, ["check-ats-structure", "--resume", str(resume_json)]
    )
    assert structure.exit_code == 0, structure.stdout
    assert json.loads(structure.stdout)["errors"] == []


def test_gate_blocks_lossy_conversion(tmp_path: Path) -> None:
    _init_project(tmp_path)
    src = tmp_path / "resume-a.txt"
    src.write_text(_SOURCE_TEXT, encoding="utf-8")

    # Drop a whole bullet: count mismatch + a multi-word dropped span → error.
    lossy = _faithful_resume()
    lossy.workExperience[0].description = ["Built Python services for analytics"]
    lossy_json = tmp_path / "lossy.json"
    lossy_json.write_text(lossy.model_dump_json(), encoding="utf-8")

    gate = runner.invoke(
        app, ["validate-faithfulness", "--source", str(src), "--json", str(lossy_json)]
    )
    assert gate.exit_code != 0, gate.stdout
    report = json.loads(gate.stdout)["data"]
    assert report["passed"] is False
    codes = {f["code"] for f in report["findings"]}
    assert "bullet_count_mismatch" in {c.lower() for c in codes}


def test_job_pointer_path(tmp_path: Path) -> None:
    _init_project(tmp_path)
    rk = tmp_path / "resume-kit"

    job_src = tmp_path / "job.txt"
    job_src.write_text(
        "Backend Engineer at Acme\nRequirements: Python, SQL, AWS\n", encoding="utf-8"
    )
    extracted = runner.invoke(app, ["extract-text", str(job_src)])
    assert extracted.exit_code == 0, extracted.stdout
    assert "Backend Engineer" in json.loads(extracted.stdout)["data"]["text"]

    # A structured JobDescription would be written by the interpretation step;
    # here we assert the pointer contract records job + source deterministically.
    job_json = rk / "jobs" / "job-original.json"
    job_json.write_text('{"raw_text": "Backend Engineer at Acme"}', encoding="utf-8")
    activate = runner.invoke(
        app,
        [
            "set-active",
            "--root",
            str(tmp_path),
            "--job",
            "jobs/job-original.json",
            "--job-source",
            str(job_src),
        ],
    )
    assert activate.exit_code == 0, activate.stdout
    config = json.loads((rk / "config.json").read_text(encoding="utf-8"))
    assert config["active_job"] == "jobs/job-original.json"
    assert config["active_job_source"] == str(job_src)


def test_deterministic_docx_extraction_no_extras(tmp_path: Path) -> None:
    if not _docx_available():
        import pytest

        pytest.skip("python-docx not installed")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Ada Lovelace")
    doc.add_paragraph("Senior Engineer")
    doc.add_paragraph("Built Python services for analytics")
    docx_path = tmp_path / "resume-a.docx"
    doc.save(str(docx_path))

    extracted = runner.invoke(app, ["extract-text", str(docx_path)])
    assert extracted.exit_code == 0, extracted.stdout
    data = json.loads(extracted.stdout)["data"]
    assert "Ada Lovelace" in data["text"]
    assert "Python" in data["text"]
