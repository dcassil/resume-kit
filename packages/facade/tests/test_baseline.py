"""original -> base build path with the claim-preservation gate (RIT-T-0115)."""

from __future__ import annotations

import json
from pathlib import Path

from resume_kit_facade.baseline import build_base
from resume_kit_facade.project_config import (
    init_project,
    load_config,
    resolve_active_resume,
    set_active,
    working_dir,
)
from resume_kit_schemas import ResumeDocument


def _setup(root: Path, resume: dict) -> None:
    init_project(root)
    (working_dir(root) / "resumes" / "jane-original.json").write_text(
        json.dumps(resume), encoding="utf-8"
    )
    set_active(root, resume="resumes/jane-original.json")


def _resume_with_pii() -> dict:
    return {
        "personalInfo": {"name": "Jane", "email": "j@x.com", "phone": "555"},
        "summary": "Engineer. SSN 123-45-6789.",
        "workExperience": [
            {"title": "Staff Engineer", "company": "Acme", "years": "Jan 2020 - Present",
             "description": ["Built billing."]},
            {
                "title": "Engineer",
                "company": "Beta",
                "years": "2016 - 2019",
                "description": ["Shipped API."],
            },
        ],
        "education": [{"institution": "MIT", "degree": "BS CS", "years": "2016"}],
        "additional": {"technicalSkills": ["Python"]},
    }


def test_build_base_writes_base_and_sets_pointer(tmp_path: Path) -> None:
    _setup(tmp_path, _resume_with_pii())
    result = build_base(tmp_path, mode="auto")

    assert result.base_path == "resumes/jane-base.json"
    assert "PII_SSN" in result.applied
    assert "INCONSISTENT_DATES" in result.applied

    base_file = working_dir(tmp_path) / "resumes" / "jane-base.json"
    assert base_file.exists()
    base = ResumeDocument.model_validate(json.loads(base_file.read_text(encoding="utf-8")))
    # PII stripped, dates normalized, claims preserved
    assert "123-45-6789" not in base.summary
    assert "Jan" not in base.workExperience[0].years
    assert base.workExperience[0].company == "Acme"

    # config points base at the new artifact; resolution prefers it
    config = load_config(tmp_path)
    assert config.base_resume == "resumes/jane-base.json"
    assert config.base_derived_from == "resumes/jane-original.json"
    assert resolve_active_resume(config) == "resumes/jane-base.json"


def test_build_base_defers_needs_judgment(tmp_path: Path) -> None:
    resume = _resume_with_pii()
    resume["personalInfo"] = {"name": "Jane", "phone": "555"}  # missing email -> needs_judgment
    resume["customSections"] = {"My Superpowers": {"sectionType": "text", "text": "x"}}
    _setup(tmp_path, resume)
    result = build_base(tmp_path, mode="auto")
    assert "MISSING_EMAIL" in result.deferred
    assert "NONSTANDARD_SECTION" in result.deferred


def test_clean_resume_base_equals_original_content(tmp_path: Path) -> None:
    clean = {
        "personalInfo": {"name": "Jane", "email": "j@x.com", "phone": "555"},
        "summary": "Engineer.",
        "workExperience": [
            {"title": "A", "company": "X", "years": "2020 - 2022", "description": ["Built."]}
        ],
        "education": [{"institution": "M", "degree": "BS", "years": "2016"}],
        "additional": {"technicalSkills": ["Python"]},
    }
    _setup(tmp_path, clean)
    result = build_base(tmp_path, mode="auto")
    assert result.applied == []
    base = json.loads((working_dir(tmp_path) / "resumes" / "jane-base.json").read_text())
    assert ResumeDocument.model_validate(base) == ResumeDocument.model_validate(clean)


# --- base -> standard walkthrough write path (RIT-T-0118) ---


def _resume_for_standard() -> dict:
    return {
        "personalInfo": {"name": "Jane", "email": "j@x.com", "phone": "5"},
        "summary": "A results-driven team player who ships.",
        "workExperience": [
            {"id": 1, "title": "Eng", "company": "Acme", "years": "2020-2022",
             "description": ["Responsible for maintaining the billing service."]}
        ],
        "additional": {"technicalSkills": ["Python"]},
    }


def test_build_standard_writes_standard_and_sets_pointer(tmp_path: Path) -> None:
    from resume_kit_facade.baseline import build_standard

    _setup(tmp_path, _resume_for_standard())
    # base first (auto), then standard from base
    build_base(tmp_path, mode="auto")
    result = build_standard(tmp_path)

    assert result.standard_path == "resumes/jane-standard.json"
    std_file = working_dir(tmp_path) / "resumes" / "jane-standard.json"
    assert std_file.exists()
    std = ResumeDocument.model_validate(json.loads(std_file.read_text(encoding="utf-8")))
    # buzzwords removed, weak opener fixed; claims preserved
    assert "results-driven" not in std.summary.lower()
    assert not std.workExperience[0].description[0].lower().startswith("responsible for")
    assert std.workExperience[0].company == "Acme"

    config = load_config(tmp_path)
    assert config.standard_resume == "resumes/jane-standard.json"
    assert config.standard_derived_from == "resumes/jane-base.json"
    assert resolve_active_resume(config) == "resumes/jane-standard.json"


def test_build_standard_applies_user_answer(tmp_path: Path) -> None:
    from datetime import date

    from resume_kit_facade.baseline import build_standard
    from resume_kit_scoring import (
        analyze_best_practices,
        finding_key,
        project_scoredoc,
    )

    _setup(tmp_path, _resume_for_standard())
    base = build_base(tmp_path, mode="auto")
    base_doc = ResumeDocument.model_validate(
        json.loads((working_dir(tmp_path) / base.base_path).read_text())
    )
    scoredoc = project_scoredoc(base_doc, reference_date=date(2025, 1, 1))
    report = analyze_best_practices(base_doc, scoredoc)
    mq = next(f for f in report.findings if f.rule_code == "MISSING_QUANTIFICATION")
    answers = {finding_key(mq): "Cut billing incidents 40% over two quarters."}

    result = build_standard(tmp_path, answers=answers)
    std = ResumeDocument.model_validate(
        json.loads((working_dir(tmp_path) / result.standard_path).read_text())
    )
    assert any("40%" in b for b in std.workExperience[0].description)


def test_build_base_surfaces_source_parse_risk(tmp_path: Path) -> None:
    """A source .docx with a table surfaces SOURCE_LAYOUT_TABLE in deferred (RIT-T-0122)."""
    import io

    from docx import Document

    _setup(tmp_path, _resume_with_pii())

    doc = Document()
    doc.add_paragraph("Jane")
    doc.add_table(rows=2, cols=2)
    buf = io.BytesIO()
    doc.save(buf)
    source_path = working_dir(tmp_path) / "resumes" / "jane-original.docx"
    source_path.write_bytes(buf.getvalue())

    set_active(
        tmp_path,
        resume="resumes/jane-original.json",
        resume_source="resumes/jane-original.docx",
    )

    result = build_base(tmp_path, mode="auto")

    # Warning surfaces as a deferred code but never gates the write.
    assert "SOURCE_LAYOUT_TABLE" in result.deferred
    assert (working_dir(tmp_path) / result.base_path).exists()


def test_build_base_no_source_has_no_risk_codes(tmp_path: Path) -> None:
    """With no source file set, build_base behaves exactly as before (RIT-T-0122)."""
    _setup(tmp_path, _resume_with_pii())
    result = build_base(tmp_path, mode="auto")
    assert not any(code.startswith("SOURCE_LAYOUT_") for code in result.deferred)
