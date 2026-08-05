"""No-LLM deterministic-path tests: no provider, no network, exit 0."""

from __future__ import annotations

import importlib.util
import io
import json
from pathlib import Path

import pytest
from resume_kit_cli.app import app
from resume_kit_schemas import JobDescription, ResumeDocument
from typer.testing import CliRunner

runner = CliRunner()


def _docx_available() -> bool:
    """Return True when python-docx is importable (optional test dep)."""
    return importlib.util.find_spec("docx") is not None


def _resume(path: Path) -> str:
    path.write_text(ResumeDocument().model_dump_json(), encoding="utf-8")
    return str(path)


def _job(path: Path) -> str:
    job = JobDescription(raw_text="Backend role. Requirements: Python, SQL, AWS.")
    path.write_text(job.model_dump_json(), encoding="utf-8")
    return str(path)


def test_extract_no_llm(tmp_path: Path) -> None:
    src = tmp_path / "resume.txt"
    src.write_text("Ada Lovelace\nEngineer\nSkills: Python\n", encoding="utf-8")
    result = runner.invoke(app, ["extract", str(src), "--no-llm"])
    assert result.exit_code == 0, result.stdout
    assert json.loads(result.stdout)["errors"] == []


def test_extract_job_no_llm(tmp_path: Path) -> None:
    src = tmp_path / "job.txt"
    src.write_text("Data Engineer. Requirements: Spark.", encoding="utf-8")
    result = runner.invoke(app, ["extract-job", str(src), "--no-llm"])
    assert result.exit_code == 0, result.stdout


def test_extract_text_txt(tmp_path: Path) -> None:
    src = tmp_path / "resume.txt"
    src.write_text("Ada Lovelace\nEngineer\nSkills: Python\n", encoding="utf-8")
    result = runner.invoke(app, ["extract-text", str(src)])
    assert result.exit_code == 0, result.stdout
    payload = json.loads(result.stdout)
    assert payload["errors"] == []
    assert "Ada Lovelace" in payload["data"]["text"]
    assert payload["data"]["extraction_method"] == "decode"


def test_extract_text_from_stdin() -> None:
    result = runner.invoke(app, ["extract-text", "-"], input="Grace Hopper\nEng\n")
    assert result.exit_code == 0, result.stdout
    payload = json.loads(result.stdout)
    assert "Grace Hopper" in payload["data"]["text"]


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
def test_extract_text_docx(tmp_path: Path) -> None:
    """TC-001: deterministic docx extraction via the CLI (base wheel only)."""
    import docx  # noqa: PLC0415

    doc = docx.Document()
    doc.add_paragraph("Marie Curie")
    doc.add_paragraph("Research Scientist")
    buf = io.BytesIO()
    doc.save(buf)
    src = tmp_path / "resume.docx"
    src.write_bytes(buf.getvalue())

    result = runner.invoke(app, ["extract-text", str(src)])
    assert result.exit_code == 0, result.stdout
    payload = json.loads(result.stdout)
    assert payload["errors"] == []
    assert "Marie Curie" in payload["data"]["text"]
    assert payload["data"]["extraction_method"] == "markitdown"


def test_extract_from_stdin(tmp_path: Path) -> None:
    result = runner.invoke(app, ["extract", "-", "--no-llm"], input="Grace Hopper\nEng\n")
    assert result.exit_code == 0, result.stdout


def test_ats_and_match_and_select(tmp_path: Path) -> None:
    resume = _resume(tmp_path / "r.json")
    job = _job(tmp_path / "j.json")
    resumes = tmp_path / "rs.json"
    resumes.write_text(
        json.dumps([ResumeDocument().model_dump(mode="json")]), encoding="utf-8"
    )
    for args in (
        ["check-ats", "--resume", resume, "--job", job],
        ["match", "--resume", resume, "--job", job],
        ["select", "--resumes", str(resumes), "--job", job],
    ):
        result = runner.invoke(app, args)
        assert result.exit_code == 0, (args, result.stdout)


def test_export_no_provider(tmp_path: Path) -> None:
    resume = _resume(tmp_path / "r.json")
    out = tmp_path / "resume.pdf"
    result = runner.invoke(
        app, ["export", "--resume", resume, "--format", "pdf", "--out", str(out)]
    )
    assert result.exit_code == 0, result.stdout
    assert out.read_bytes().startswith(b"%PDF-")


def test_compare_gaps_truth_evidence(tmp_path: Path) -> None:
    resume = _resume(tmp_path / "r.json")
    job = _job(tmp_path / "j.json")
    for args in (
        ["compare", "--base", resume, "--candidate", resume, "--job", job],
        ["identify-gaps", "--job", job, "--tailored", resume, "--master", resume],
        ["validate-truth", "--resume", resume],
        ["build-evidence", "--resume", resume],
    ):
        result = runner.invoke(app, args)
        assert result.exit_code == 0, (args, result.stdout)
