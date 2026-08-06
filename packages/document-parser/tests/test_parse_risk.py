"""Unit tests for the source parse-risk detector (RIT-T-0122)."""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from resume_kit_document_parser.parse_risk import (
    SOURCE_LAYOUT_HEADER_FOOTER,
    SOURCE_LAYOUT_IMAGE_ONLY,
    SOURCE_LAYOUT_MULTICOLUMN,
    SOURCE_LAYOUT_TABLE,
    SOURCE_LAYOUT_TEXTBOX,
    detect_source_parse_risks,
)
from resume_kit_schemas import FindingSeverity, FixAffordance


def _docx_bytes(document: Document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _codes(content: bytes, filename: str = "resume.docx") -> set[str]:
    return {f.code for f in detect_source_parse_risks(content, filename)}


def test_table_docx_emits_table_code() -> None:
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_table(rows=2, cols=2)
    findings = detect_source_parse_risks(_docx_bytes(doc), "resume.docx")
    codes = {f.code for f in findings}
    assert SOURCE_LAYOUT_TABLE in codes
    table_finding = next(f for f in findings if f.code == SOURCE_LAYOUT_TABLE)
    assert table_finding.severity is FindingSeverity.WARNING
    assert table_finding.fix_affordance is FixAffordance.NEEDS_JUDGMENT


def test_multicolumn_docx_emits_multicolumn_code() -> None:
    doc = Document()
    doc.add_paragraph("Hello")
    sect_pr = doc.sections[0]._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = sect_pr.makeelement(qn("w:cols"), {})
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    codes = _codes(_docx_bytes(doc))
    assert SOURCE_LAYOUT_MULTICOLUMN in codes


def test_header_footer_docx_emits_header_footer_code() -> None:
    doc = Document()
    doc.add_paragraph("Body")
    doc.sections[0].header.paragraphs[0].add_run("My Name — Contact")
    codes = _codes(_docx_bytes(doc))
    assert SOURCE_LAYOUT_HEADER_FOOTER in codes


def test_textbox_docx_emits_textbox_code() -> None:
    doc = Document()
    para = doc.add_paragraph("Body")
    # Inject a minimal element carrying a txbxContent tag into the paragraph.
    txbx = para._p.makeelement(qn("w:txbxContent"), {})
    para._p.append(txbx)
    assert "txbxContent" in doc.element.body.xml
    codes = _codes(_docx_bytes(doc))
    assert SOURCE_LAYOUT_TEXTBOX in codes


def test_clean_docx_returns_no_findings() -> None:
    doc = Document()
    doc.add_paragraph("A perfectly ordinary single-column paragraph.")
    content = _docx_bytes(doc)
    first = detect_source_parse_risks(content, "resume.docx")
    second = detect_source_parse_risks(content, "resume.docx")
    assert first == []
    assert first == second  # deterministic


def test_garbage_docx_bytes_never_raises() -> None:
    assert detect_source_parse_risks(b"not a docx", "resume.docx") == []


def test_unknown_extension_returns_empty() -> None:
    assert detect_source_parse_risks(b"whatever", "resume.txt") == []


def test_pdf_image_only_detected(monkeypatch: pytest.MonkeyPatch) -> None:
    import resume_kit_document_parser.parse_risk as mod

    class _Result:
        text = "   \n\t "

    monkeypatch.setattr(mod, "extract_resume_text", lambda content, filename: _Result())
    codes = _codes(b"%PDF-1.4 fake", "resume.pdf")
    assert SOURCE_LAYOUT_IMAGE_ONLY in codes


def test_pdf_with_text_no_image_only(monkeypatch: pytest.MonkeyPatch) -> None:
    import resume_kit_document_parser.parse_risk as mod

    class _Result:
        text = "Jane Doe\nSoftware Engineer\nExperience"

    monkeypatch.setattr(mod, "extract_resume_text", lambda content, filename: _Result())
    codes = _codes(b"%PDF-1.4 fake", "resume.pdf")
    assert SOURCE_LAYOUT_IMAGE_ONLY not in codes
