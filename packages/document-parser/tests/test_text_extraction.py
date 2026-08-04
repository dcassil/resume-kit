"""Tests for resume_kit_document_parser.text_extraction (no-LLM path)."""

from __future__ import annotations

import importlib
from pathlib import Path

import pytest
from resume_kit_document_parser.text_extraction import (
    TextExtractionResult,
    extract_resume_text,
    parse_document,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

FIXTURES_DIR = Path(__file__).parent / "fixtures"
SAMPLE_RESUME_MD = FIXTURES_DIR / "sample_resume.md"


def _docx_available() -> bool:
    """Return True when python-docx is importable (optional test dep)."""
    return importlib.util.find_spec("docx") is not None


# ---------------------------------------------------------------------------
# Plain-text / Markdown decode path
# ---------------------------------------------------------------------------


class TestPlainTextDecode:
    def test_md_fixture_roundtrip(self) -> None:
        """Reading the sample fixture via extract_resume_text returns its content."""
        raw = SAMPLE_RESUME_MD.read_bytes()
        result = extract_resume_text(raw, "sample_resume.md")

        assert isinstance(result, TextExtractionResult)
        assert result.extraction_method == "decode"
        assert result.filename == "sample_resume.md"
        assert result.warnings == []
        # Spot-check some content
        assert "Jane Doe" in result.text
        assert "Senior Software Engineer" in result.text

    def test_txt_extension_decode(self) -> None:
        content = b"Name: Alice\nSkills: Python, SQL\n"
        result = extract_resume_text(content, "resume.txt")

        assert result.extraction_method == "decode"
        assert result.text == "Name: Alice\nSkills: Python, SQL\n"
        assert result.warnings == []

    def test_markdown_extension(self) -> None:
        content = b"# Resume\n\n## Experience\n\nDid things.\n"
        result = extract_resume_text(content, "cv.markdown")

        assert result.extraction_method == "decode"
        assert "Experience" in result.text
        assert result.warnings == []

    def test_invalid_utf8_yields_warning_not_raise(self) -> None:
        bad_bytes = b"Name: \xff\xfe broken"
        result = extract_resume_text(bad_bytes, "bad.txt")

        assert result.extraction_method == "decode"
        # replacement char injected
        assert "�" in result.text
        assert len(result.warnings) == 1
        assert result.warnings[0].code == "unknown"


# ---------------------------------------------------------------------------
# Empty / unsupported input
# ---------------------------------------------------------------------------


class TestEdgeCases:
    def test_empty_content_yields_warning_not_raise(self) -> None:
        result = extract_resume_text(b"", "resume.pdf")

        assert result.text == ""
        assert len(result.warnings) == 1
        assert result.warnings[0].code == "unknown"
        assert "Empty" in result.warnings[0].message
        assert result.extraction_method == "none"

    def test_unsupported_extension_yields_warning_not_raise(self) -> None:
        result = extract_resume_text(b"some data", "resume.odt")

        assert result.text == ""
        assert len(result.warnings) == 1
        assert result.warnings[0].code == "unknown"
        assert "Unsupported" in result.warnings[0].message
        assert result.extraction_method == "none"

    def test_no_extension_yields_warning_not_raise(self) -> None:
        result = extract_resume_text(b"plain text", "resume")

        assert result.text == ""
        assert len(result.warnings) == 1


# ---------------------------------------------------------------------------
# MarkItDown-backed path (DOCX)
# ---------------------------------------------------------------------------


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
class TestMarkItDownDocx:
    def test_simple_docx_roundtrip(self) -> None:
        """Build a minimal DOCX in-memory and verify MarkItDown extracts text."""
        import io

        import docx  # noqa: PLC0415

        doc = docx.Document()
        doc.add_heading("Jane Doe", level=1)
        doc.add_paragraph("Senior Software Engineer")
        doc.add_paragraph("Python, Go, Kubernetes")

        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = extract_resume_text(content, "resume.docx")

        assert result.extraction_method == "markitdown"
        assert result.filename == "resume.docx"
        assert result.warnings == []
        assert "Jane Doe" in result.text or "Jane" in result.text

    def test_parse_document_returns_string(self) -> None:
        """parse_document (low-level) returns a str for a valid DOCX."""
        import io

        import docx  # noqa: PLC0415

        doc = docx.Document()
        doc.add_paragraph("Hello from parse_document")
        buf = io.BytesIO()
        doc.save(buf)

        text = parse_document(buf.getvalue(), "test.docx")

        assert isinstance(text, str)
        assert len(text) > 0


# ---------------------------------------------------------------------------
# Corrupt DOCX yields warning, not raise
# ---------------------------------------------------------------------------


class TestMarkItDownFailure:
    def test_corrupt_docx_does_not_raise(self) -> None:
        """Garbage bytes with .docx extension must never raise.

        MarkItDown may succeed (returning garbage text) or raise internally —
        either way ``extract_resume_text`` must return a ``TextExtractionResult``
        with ``extraction_method == "markitdown"``.  If conversion fails a
        warning is appended; if MarkItDown surprisingly succeeds the warnings
        list may be empty.
        """
        result = extract_resume_text(b"not a real docx file at all", "corrupt.docx")

        assert isinstance(result, TextExtractionResult)
        assert result.extraction_method == "markitdown"
        assert result.filename == "corrupt.docx"
        # Regardless of success/failure, the call must not raise.
