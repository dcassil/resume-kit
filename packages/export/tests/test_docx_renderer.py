"""Tests for deterministic DOCX rendering."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from resume_kit_export.docx import render_docx
from resume_kit_schemas.resume import ResumeDocument, SectionType


def test_render_docx_returns_reopenable_bytes() -> None:
    data = render_docx(_resume())

    assert isinstance(data, bytes)
    assert len(data) > 0
    document = Document(BytesIO(data))
    assert _paragraph_text(document)[0] == "Jane Doe"


def test_render_docx_is_byte_deterministic() -> None:
    resume = _resume()

    first = render_docx(resume)
    second = render_docx(resume)

    # The renderer fixes core properties and rewrites ZIP entries, so raw bytes
    # are stable. This is stronger than a structural-fingerprint comparison.
    assert first == second


def test_render_docx_structural_content_order() -> None:
    document = Document(BytesIO(render_docx(_resume())))
    paragraphs = _paragraph_text(document)

    _assert_in_order(
        paragraphs,
        [
            "Jane Doe",
            "Principal Engineer",
            (
                "jane@example.com | 555-0100 | Chicago, IL | https://jane.example | "
                "linkedin.com/in/janedoe | github.com/janedoe"
            ),
            "Summary",
            "Builds deterministic document systems.",
            "Experience",
            "Principal Engineer | Example Co | Remote | 2020-2026",
            "Led export platform migration.",
            "Education",
            "MS Computer Science | Example University | 2016-2018",
            "Projects",
            "Resume Kit | Maintainer | 2024-2026 | github.com/example/resume-kit",
            "Built schema-driven exporters.",
            "Skills & Awards",
            "Technical Skills: Python, Pydantic, DOCX",
            "Languages: English, Spanish",
            "Certifications & Training: Document Systems Certificate",
            "Awards: Reliability Award",
            "Writing",
            "Technical publishing and release-note editing.",
        ],
    )
    assert "Hidden Notes" not in paragraphs


def test_render_docx_omits_empty_optional_fields_and_placeholders() -> None:
    data = render_docx(ResumeDocument())
    document = Document(BytesIO(data))
    paragraphs = _paragraph_text(document)

    assert paragraphs == []
    assert "None" not in paragraphs
    assert "N/A" not in paragraphs


def _resume() -> ResumeDocument:
    return ResumeDocument(
        personalInfo={
            "name": "Jane Doe",
            "title": "Principal Engineer",
            "email": "jane@example.com",
            "phone": "555-0100",
            "location": "Chicago, IL",
            "website": "https://jane.example",
            "linkedin": "linkedin.com/in/janedoe",
            "github": "github.com/janedoe",
        },
        summary="Builds deterministic document systems.",
        workExperience=[
            {
                "title": "Principal Engineer",
                "company": "Example Co",
                "location": "Remote",
                "years": "2020-2026",
                "description": ["Led export platform migration."],
            }
        ],
        education=[
            {
                "institution": "Example University",
                "degree": "MS Computer Science",
                "years": "2016-2018",
            }
        ],
        personalProjects=[
            {
                "name": "Resume Kit",
                "role": "Maintainer",
                "years": "2024-2026",
                "github": "github.com/example/resume-kit",
                "description": ["Built schema-driven exporters."],
            }
        ],
        additional={
            "technicalSkills": ["Python", "Pydantic", "DOCX"],
            "languages": ["English", "Spanish"],
            "certificationsTraining": ["Document Systems Certificate"],
            "awards": ["Reliability Award"],
        },
        sectionMeta=[
            {
                "id": "writing",
                "key": "writing",
                "displayName": "Writing",
                "sectionType": SectionType.TEXT,
                "isDefault": False,
                "isVisible": True,
                "order": 6,
            },
            {
                "id": "hiddenNotes",
                "key": "hiddenNotes",
                "displayName": "Hidden Notes",
                "sectionType": SectionType.TEXT,
                "isDefault": False,
                "isVisible": False,
                "order": 7,
            },
        ],
        customSections={
            "hiddenNotes": {
                "sectionType": SectionType.TEXT,
                "text": "This should not render.",
            },
            "writing": {
                "sectionType": SectionType.TEXT,
                "text": "Technical publishing and release-note editing.",
            },
        },
    )


def _paragraph_text(document: object) -> list[str]:
    return [
        paragraph.text
        for paragraph in document.paragraphs  # type: ignore[attr-defined]
        if paragraph.text
    ]


def _assert_in_order(paragraphs: list[str], expected: list[str]) -> None:
    position = 0
    for expected_text in expected:
        for index in range(position, len(paragraphs)):
            if expected_text in paragraphs[index]:
                position = index + 1
                break
        else:
            raise AssertionError(f"{expected_text!r} not found after index {position}")
