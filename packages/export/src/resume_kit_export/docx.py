"""Deterministic DOCX rendering for structured resumes."""

from __future__ import annotations

from collections.abc import Iterable, Sequence
from datetime import datetime
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.shared import Mm, Pt
from resume_kit_schemas.resume import ResumeDocument, SectionType

from resume_kit_export.dedupe import dedupe_skill_custom_sections_for_render
from resume_kit_export.models import ExportOptions

_FIXED_DATETIME = datetime(2000, 1, 1, 0, 0, 0)
_FIXED_ZIP_DATE_TIME = (2000, 1, 1, 0, 0, 0)
_DOCX_AUTHOR = "resume-kit"
_DOCX_TITLE = "resume-kit DOCX export"


def render_docx(
    resume: ResumeDocument, *, options: ExportOptions | None = None
) -> bytes:
    """Render *resume* as deterministic DOCX bytes."""

    resume = dedupe_skill_custom_sections_for_render(resume)
    resolved_options = options or ExportOptions()
    document = Document()
    _configure_document(document, resolved_options)
    _set_core_properties(document)

    _render_header(document, resume)
    _render_summary(document, resume)
    _render_experience(document, resume)
    _render_education(document, resume)
    _render_projects(document, resume)
    _render_additional(document, resume)
    _render_custom_sections(document, resume)

    raw = BytesIO()
    document.save(raw)
    return _normalize_zip(raw.getvalue())


def _configure_document(document: DocxDocument, options: ExportOptions) -> None:
    margin = Mm(options.margin_mm)
    for section in document.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = margin
        section.right_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin

    normal_style = document.styles["Normal"]
    normal_style.font.name = options.font_family
    normal_style.font.size = Pt(options.font_size_pt)


def _set_core_properties(document: DocxDocument) -> None:
    properties = document.core_properties
    properties.author = _DOCX_AUTHOR
    properties.category = "resume"
    properties.comments = ""
    properties.content_status = "final"
    properties.created = _FIXED_DATETIME
    properties.identifier = "resume-kit-docx"
    properties.keywords = "resume-kit"
    properties.language = "en-US"
    properties.last_modified_by = _DOCX_AUTHOR
    properties.modified = _FIXED_DATETIME
    properties.revision = 1
    properties.subject = "deterministic resume export"
    properties.title = _DOCX_TITLE
    properties.version = "1"


def _render_header(document: DocxDocument, resume: ResumeDocument) -> None:
    personal = resume.personalInfo
    if name := _clean(personal.name):
        document.add_heading(name, level=0)
    if title := _clean(personal.title):
        document.add_paragraph(title)

    contact = _clean_join(
        [
            personal.email,
            personal.phone,
            personal.location,
            personal.website,
            personal.linkedin,
            personal.github,
        ],
        separator=" | ",
    )
    if contact:
        document.add_paragraph(contact)


def _render_summary(document: DocxDocument, resume: ResumeDocument) -> None:
    if summary := _clean(resume.summary):
        document.add_heading("Summary", level=1)
        document.add_paragraph(summary)


def _render_experience(document: DocxDocument, resume: ResumeDocument) -> None:
    entries = [
        entry
        for entry in resume.workExperience
        if _has_text(
            [entry.title, entry.company, entry.location, entry.years, *entry.description]
        )
    ]
    if not entries:
        return

    document.add_heading("Experience", level=1)
    for entry in entries:
        if line := _clean_join(
            [entry.title, entry.company, entry.location, entry.years], separator=" | "
        ):
            document.add_paragraph(line)
        _add_bullets(document, entry.description)


def _render_education(document: DocxDocument, resume: ResumeDocument) -> None:
    entries = [
        entry
        for entry in resume.education
        if _has_text([entry.institution, entry.degree, entry.years, entry.description])
    ]
    if not entries:
        return

    document.add_heading("Education", level=1)
    for entry in entries:
        if line := _clean_join(
            [entry.degree, entry.institution, entry.years], separator=" | "
        ):
            document.add_paragraph(line)
        if description := _clean(entry.description):
            document.add_paragraph(description)


def _render_projects(document: DocxDocument, resume: ResumeDocument) -> None:
    entries = [
        entry
        for entry in resume.personalProjects
        if _has_text(
            [
                entry.name,
                entry.role,
                entry.years,
                entry.github,
                entry.website,
                *entry.description,
            ]
        )
    ]
    if not entries:
        return

    document.add_heading("Projects", level=1)
    for entry in entries:
        if line := _clean_join(
            [entry.name, entry.role, entry.years, entry.github, entry.website],
            separator=" | ",
        ):
            document.add_paragraph(line)
        _add_bullets(document, entry.description)


def _render_additional(document: DocxDocument, resume: ResumeDocument) -> None:
    additional = resume.additional
    rows = [
        ("Technical Skills", additional.technicalSkills),
        ("Languages", additional.languages),
        ("Certifications & Training", additional.certificationsTraining),
        ("Awards", additional.awards),
    ]
    populated_rows = [
        (label, values) for label, values in rows if _clean_join(values, separator=", ")
    ]
    if not populated_rows:
        return

    document.add_heading("Skills & Awards", level=1)
    for label, values in populated_rows:
        document.add_paragraph(f"{label}: {_clean_join(values, separator=', ')}")


def _render_custom_sections(document: DocxDocument, resume: ResumeDocument) -> None:
    meta_by_key = {meta.key: meta for meta in resume.sectionMeta}
    meta_by_id = {meta.id: meta for meta in resume.sectionMeta}

    def sort_key(key: str) -> tuple[int, str, str]:
        meta = meta_by_key.get(key) or meta_by_id.get(key)
        display_name = meta.displayName if meta else key
        order = meta.order if meta else 10_000
        return (order, display_name, key)

    for key in sorted(resume.customSections, key=sort_key):
        section = resume.customSections[key]
        meta = meta_by_key.get(key) or meta_by_id.get(key)
        if meta and not meta.isVisible:
            continue

        display_name = meta.displayName if meta else key
        if section.sectionType == SectionType.TEXT:
            if text := _clean(section.text):
                document.add_heading(display_name, level=1)
                document.add_paragraph(text)
        elif section.sectionType == SectionType.STRING_LIST:
            values = _clean_values(section.strings or [])
            if values:
                document.add_heading(display_name, level=1)
                _add_bullets(document, values)
        elif section.sectionType == SectionType.ITEM_LIST:
            items = section.items or []
            visible_items = [
                item
                for item in items
                if _has_text(
                    [
                        item.title,
                        item.subtitle,
                        item.location,
                        item.years,
                        *item.description,
                    ]
                )
            ]
            if visible_items:
                document.add_heading(display_name, level=1)
                for item in visible_items:
                    if line := _clean_join(
                        [item.title, item.subtitle, item.location, item.years],
                        separator=" | ",
                    ):
                        document.add_paragraph(line)
                    _add_bullets(document, item.description)


def _add_bullets(document: DocxDocument, values: Iterable[str | None]) -> None:
    for value in _clean_values(values):
        document.add_paragraph(value, style="List Bullet")


def _clean(value: str | None) -> str:
    return value.strip() if value else ""


def _clean_values(values: Iterable[str | None]) -> list[str]:
    return [cleaned for value in values if (cleaned := _clean(value))]


def _clean_join(values: Iterable[str | None], *, separator: str) -> str:
    return separator.join(_clean_values(values))


def _has_text(values: Sequence[str | None]) -> bool:
    return any(_clean(value) for value in values)


def _normalize_zip(data: bytes) -> bytes:
    source_buffer = BytesIO(data)
    target_buffer = BytesIO()
    with ZipFile(source_buffer, "r") as source, ZipFile(
        target_buffer,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        target.comment = b""
        for name in sorted(source.namelist()):
            payload = source.read(name)
            source_info = source.getinfo(name)
            target_info = ZipInfo(filename=name, date_time=_FIXED_ZIP_DATE_TIME)
            target_info.compress_type = ZIP_DEFLATED
            target_info.external_attr = source_info.external_attr
            target_info.internal_attr = source_info.internal_attr
            target_info.create_system = 0
            target.writestr(target_info, payload)
    return target_buffer.getvalue()
